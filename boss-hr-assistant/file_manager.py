"""文件管理模块 — 简历重命名、筛选、归档.

功能：
- 根据95分阈值筛选优质简历
- 按固定格式重命名简历文件
- 将优质简历复制到 shortlisted 目录
- 生成筛选报告
"""
import logging
import os
import re
import shutil
from datetime import datetime
from pathlib import Path

from config import RESUME_DIR, SHORTLISTED_DIR, SCORE_THRESHOLD
import config as _config
from resume_evaluator import ResumeEvaluator

log = logging.getLogger("boss-file-manager")

os.makedirs(SHORTLISTED_DIR, exist_ok=True)


class FileManager:
    """管理简历文件的重命名、筛选和归档.

    动态读取 config.PROFILE，支持运行时切换岗位。
    """

    def __init__(self):
        self.evaluator = ResumeEvaluator()
        self.threshold = SCORE_THRESHOLD

    @property
    def naming_format(self) -> str:
        return _config.PROFILE.get("naming", {}).get(
            "format", "{name}_{job_title}_{score}分_{city}_{experience}"
        )

    @property
    def job_config(self) -> dict:
        return _config.PROFILE.get("job", {})

    def process_downloaded_resumes(self, resume_dir: str = None,
                                   auto_match: bool = False,
                                   chat_list_info: list[dict] = None) -> dict:
        """处理已下载的简历文件：评分、筛选、重命名.

        Args:
            resume_dir: 简历目录，默认使用配置目录
            auto_match: 是否自动根据简历内容匹配岗位（当未指定固定岗位时）
            chat_list_info: 聊天列表信息（含候选人投递职位），用于按候选人匹配岗位
        Returns:
            处理结果统计
        """
        target_dir = resume_dir or RESUME_DIR
        if not os.path.exists(target_dir):
            return {"error": f"简历目录不存在: {target_dir}"}

        # 扫描所有简历文件
        resume_files = self._scan_resume_files(target_dir)
        if not resume_files:
            return {"error": "未找到简历文件"}

        # 构建候选人→职位映射（如果提供了聊天列表）
        candidate_job_map = {}
        if chat_list_info:
            for chat in chat_list_info:
                if isinstance(chat, dict) and "error" not in chat:
                    name = chat.get("name", "").strip()
                    job_title = chat.get("job_title", "").strip()
                    if name and job_title:
                        candidate_job_map[name] = job_title

        results = []
        shortlisted = []

        for file_path in resume_files:
            result = self._process_single_file(
                file_path,
                auto_match=auto_match,
                candidate_job_map=candidate_job_map,
            )
            results.append(result)

            if result.get("passed"):
                shortlisted.append(result)

        # 生成筛选报告
        report = self._generate_report(results, shortlisted)

        return {
            "status": "success",
            "total": len(results),
            "shortlisted": len(shortlisted),
            "rejected": len(results) - len(shortlisted),
            "shortlisted_files": [r["new_path"] for r in shortlisted if r.get("new_path")],
            "report": report,
            "details": results,
        }

    def _scan_resume_files(self, directory: str) -> list[str]:
        """扫描目录中的简历文件."""
        extensions = {".pdf", ".doc", ".docx", ".txt"}
        files = []

        for root, dirs, filenames in os.walk(directory):
            # 跳过 shortlisted 子目录（避免重复处理）
            if "shortlisted" in root:
                continue
            for filename in filenames:
                ext = Path(filename).suffix.lower()
                if ext in extensions:
                    files.append(os.path.join(root, filename))

        return sorted(files)

    def _process_single_file(self, file_path: str,
                              auto_match: bool = False,
                              candidate_job_map: dict = None) -> dict:
        """处理单个简历文件：提取文本 → 自动匹配岗位 → 评分 → 重命名 → 归档."""
        filename = os.path.basename(file_path)
        ext = Path(file_path).suffix

        # 提取简历文本
        resume_text = self._extract_text_sync(file_path)

        # 提取候选人信息
        name = self.evaluator.extract_candidate_name(resume_text)
        experience = self.evaluator.extract_experience_years(resume_text)
        education = self.evaluator.extract_education(resume_text)

        # === 自动岗位匹配 ===
        matched_profile = None
        matched_job_title = ""
        matched_confidence = ""
        used_profile = None

        # 策略1: 从候选人→职位映射查找（来自聊天列表）
        if candidate_job_map and name in candidate_job_map:
            cand_job_title = candidate_job_map[name]
            match_result = _config.match_job_for_candidate(cand_job_title)
            if match_result:
                matched_profile = match_result
                matched_job_title = match_result.get("title", "")
                matched_confidence = match_result.get("confidence", "")
                used_profile = match_result.get("profile")
                log.info(f"[{name}] 聊天列表职位 '{cand_job_title}' → 匹配: {matched_job_title} ({matched_confidence})")

        # 策略2: 如果策略1没匹配上，尝试从简历文本中提取职位并匹配
        if not matched_profile and auto_match:
            match_result = _config.match_job_by_keywords(resume_text[:500])
            if match_result:
                matched_profile = match_result
                matched_job_title = match_result.get("title", "")
                matched_confidence = match_result.get("confidence", "")
                used_profile = match_result.get("profile")
                log.info(f"[{name}] 简历关键词匹配 → {matched_job_title} ({matched_confidence})")

        # 评分（传入匹配到的 profile，或使用默认配置）
        evaluation = self.evaluator.evaluate(
            resume_text,
            profile=used_profile,
        )
        score = evaluation["score"]

        # 确定文件名中使用的岗位标题和城市
        job_title_for_name = matched_job_title or self.job_config.get("title", "未知岗位")
        city_for_name = (matched_profile.get("city", "") if matched_profile
                         else self.job_config.get("city", ""))

        # 生成新文件名
        new_filename = self._generate_filename(
            name=name,
            job_title=job_title_for_name,
            score=score,
            city=city_for_name,
            experience=experience,
        )
        new_filename += ext

        # 重命名文件
        new_path = os.path.join(os.path.dirname(file_path), new_filename)
        if file_path != new_path:
            # 避免文件名冲突
            counter = 1
            while os.path.exists(new_path):
                base = Path(new_filename).stem
                new_filename = f"{base}_{counter}{ext}"
                new_path = os.path.join(os.path.dirname(file_path), new_filename)
            os.rename(file_path, new_path)

        result = {
            "original_name": filename,
            "new_name": new_filename,
            "new_path": new_path,
            "name": name,
            "experience": experience,
            "education": education,
            "score": score,
            "level": evaluation["level"],
            "passed": evaluation["passed"],
            "matched_job": matched_job_title or evaluation.get("matched_job", ""),
            "match_confidence": matched_confidence,
            "strengths": evaluation["strengths"],
            "weaknesses": evaluation["weaknesses"],
            "summary": evaluation["summary"],
        }

        # 如果通过筛选，复制到 shortlisted 目录
        if evaluation["passed"]:
            shortlisted_path = os.path.join(SHORTLISTED_DIR, new_filename)
            shutil.copy2(new_path, shortlisted_path)
            result["shortlisted_path"] = shortlisted_path
            log.info(f"优质简历已归档: {new_filename} (score={score})")

        return result

    def _generate_filename(self, **kwargs) -> str:
        """根据配置格式生成文件名."""
        # 替换变量
        name = kwargs.get("name", "未知")
        # 清理姓名中的特殊字符
        name = re.sub(r'[\\/:*?"<>|]', '_', name)

        params = {
            "name": name,
            "job_title": kwargs.get("job_title", ""),
            "score": kwargs.get("score", 0),
            "city": kwargs.get("city", ""),
            "experience": kwargs.get("experience", ""),
            "date": datetime.now().strftime("%Y%m%d"),
        }

        try:
            filename = self.naming_format.format(**params)
        except KeyError as e:
            log.warning(f"命名格式变量缺失: {e}")
            filename = f"{name}_{params['score']}分_{params['date']}"

        # 清理文件名
        filename = re.sub(r'[\\/:*?"<>|]', '_', filename)
        return filename

    def _extract_text_sync(self, file_path: str) -> str:
        """同步提取简历文本."""
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".pdf":
                return self._extract_pdf_sync(file_path)
            elif ext in (".doc", ".docx"):
                return self._extract_docx_sync(file_path)
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
        except Exception as e:
            log.warning(f"文本提取失败 ({ext}): {e}")
            return ""

    def _extract_pdf_sync(self, file_path: str) -> str:
        """同步提取PDF文本."""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            log.warning(f"PDF解析失败: {e}")
            return ""

    def _extract_docx_sync(self, file_path: str) -> str:
        """同步提取Word文本."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            log.warning(f"Word解析失败: {e}")
            return ""

    def _generate_report(self, all_results: list[dict], shortlisted: list[dict]) -> str:
        """生成 Markdown 格式的筛选报告."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        lines = [
            f"# 简历筛选报告",
            f"",
            f"生成时间: {timestamp}",
            f"岗位: {self.job_config.get('title', '未指定')}",
            f"城市: {self.job_config.get('city', '未指定')}",
            f"筛选阈值: {self.threshold}分",
            f"",
            f"## 统计概览",
            f"",
            f"| 项目 | 数量 |",
            f"|---|---|",
            f"| 总简历数 | {len(all_results)} |",
            f"| 通过筛选 (≥{self.threshold}分) | {len(shortlisted)} |",
            f"| 未通过 | {len(all_results) - len(shortlisted)} |",
            f"",
        ]

        if shortlisted:
            lines.extend([
                f"## 优质简历列表 (≥{self.threshold}分)",
                f"",
                f"| 姓名 | 评分 | 经验 | 学历 | 文件名 |",
                f"|---|---|---|---|---|",
            ])
            for r in sorted(shortlisted, key=lambda x: x["score"], reverse=True):
                lines.append(
                    f"| {r['name']} | {r['score']} | {r['experience']} | {r['education']} | {r['new_name']} |"
                )
            lines.append("")

        # 所有简历详情
        lines.extend([
            f"## 所有简历评分详情",
            f"",
        ])

        for r in sorted(all_results, key=lambda x: x["score"], reverse=True):
            status = "✓ 通过" if r["passed"] else "✗ 未通过"
            lines.extend([
                f"### {r['name']} — {r['score']}分 ({r['level']}) {status}",
                f"",
                f"- 文件: {r['original_name']} → {r['new_name']}",
                f"- 经验: {r['experience']} | 学历: {r['education']}",
                f"- 亮点: {'; '.join(r['strengths'][:3])}",
                f"- 风险: {'; '.join(r['weaknesses'][:3])}",
                f"",
            ])

        report = "\n".join(lines)
        report_path = os.path.join(SHORTLISTED_DIR, f"筛选报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md")

        with open(report_path, "w", encoding="utf-8") as f:
            f.write(report)

        return report_path
