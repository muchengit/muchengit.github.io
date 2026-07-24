"""简历下载模块 — 下载聊天中的简历附件并提取文本内容.

支持：
- 点击聊天中的简历附件下载
- 拦截网络请求捕获简历文件URL
- 解析 PDF/Word 简历提取文本
"""
import asyncio
import logging
import os
import re
import json
from pathlib import Path

from browser import BossBrowser
from config import RESUME_DIR

log = logging.getLogger("boss-resume-downloader")

# 确保简历目录存在
os.makedirs(RESUME_DIR, exist_ok=True)


class ResumeDownloader:
    """下载并解析 Boss 直聘聊天中的简历."""

    def __init__(self, browser: BossBrowser):
        self.browser = browser
        self._downloaded_urls: set[str] = set()

    async def download_resume_from_chat(self, chat_index: int = -1) -> dict:
        """从当前或指定聊天中下载简历.

        Args:
            chat_index: 聊天列表索引，-1 表示当前对话
        Returns:
            下载结果
        """
        from chat_scraper import ChatScraper
        scraper = ChatScraper(self.browser)

        if chat_index >= 0:
            await scraper.select_chat_by_index(chat_index)
            await asyncio.sleep(2)

        # 检查是否有简历附件
        resume_check = await scraper.check_resume_received()
        if not resume_check.get("has_resume"):
            return {"error": "当前对话未检测到简历附件"}

        results = []
        for resume in resume_check.get("resumes", []):
            result = await self._download_single_resume(resume)
            results.append(result)

        return {
            "status": "success",
            "downloaded": len([r for r in results if r.get("status") == "success"]),
            "failed": len([r for r in results if r.get("status") == "error"]),
            "details": results,
        }

    async def _download_single_resume(self, resume_info: dict) -> dict:
        """下载单个简历附件."""
        url = resume_info.get("attachment_url", "")
        text = resume_info.get("attachment_text", "")

        # 如果有直接URL，通过Playwright下载
        if url and url not in self._downloaded_urls:
            self._downloaded_urls.add(url)
            return await self._download_by_url(url, text)

        # 没有URL，尝试点击附件元素下载
        return await self._download_by_click(text)

    async def _download_by_url(self, url: str, filename_hint: str = "") -> dict:
        """通过URL下载简历文件."""
        p = self.browser.page

        try:
            # 使用Playwright的下载处理
            async with p.expect_download(timeout=30000) as download_info:
                await p.goto(url)

            download = download_info.value
            # 生成文件名
            original_name = download.suggested_filename
            if not original_name or "." not in original_name:
                original_name = (filename_hint or "resume") + ".pdf"

            # 清理文件名
            safe_name = re.sub(r'[\\/:*?"<>|]', '_', original_name)
            save_path = os.path.join(RESUME_DIR, safe_name)

            await download.save_as(save_path)

            # 提取简历文本
            resume_text = await self._extract_text(save_path)

            return {
                "status": "success",
                "file_path": save_path,
                "filename": safe_name,
                "file_size": os.path.getsize(save_path),
                "resume_text": resume_text[:3000] if resume_text else "",
                "url": url,
            }
        except Exception as e:
            log.warning(f"URL下载失败: {e}")
            # 回退: 尝试通过API下载
            return await self._download_via_api(url, filename_hint)

    async def _download_via_api(self, url: str, filename_hint: str = "") -> dict:
        """通过API请求下载简历."""
        p = self.browser.page

        try:
            # 使用页面上下文发起请求
            response = await p.request.get(url)
            if response.ok:
                content_type = response.headers.get("content-type", "")

                # 确定文件扩展名
                ext = ".pdf"
                if "word" in content_type or "docx" in url.lower():
                    ext = ".docx"
                elif "msword" in content_type or ".doc" in url.lower():
                    ext = ".doc"

                filename = (filename_hint or "resume") + ext
                safe_name = re.sub(r'[\\/:*?"<>|]', '_', filename)
                save_path = os.path.join(RESUME_DIR, safe_name)

                body = await response.body()
                with open(save_path, "wb") as f:
                    f.write(body)

                resume_text = await self._extract_text(save_path)

                return {
                    "status": "success",
                    "file_path": save_path,
                    "filename": safe_name,
                    "file_size": len(body),
                    "resume_text": resume_text[:3000] if resume_text else "",
                    "url": url,
                }
        except Exception as e:
            log.error(f"API下载失败: {e}")

        return {"status": "error", "message": f"下载失败: {url[:100]}", "url": url}

    async def _download_by_click(self, filename_hint: str = "") -> dict:
        """通过点击附件元素触发下载."""
        p = self.browser.page

        # 附件元素选择器
        attachment_selectors = [
            '[class*="file-card"]',
            '[class*="resume-card"]',
            '[class*="attach"]',
            'a[class*="file"]',
            '[class*="card-file"]',
        ]

        try:
            async with p.expect_download(timeout=15000) as download_info:
                for sel in attachment_selectors:
                    el = await p.query_selector(sel)
                    if el:
                        await el.click()
                        break
                else:
                    # iframe中查找
                    for frame in p.frames:
                        for sel in attachment_selectors:
                            try:
                                el = await frame.query_selector(sel)
                                if el:
                                    await el.click()
                                    break
                            except Exception:
                                continue
                        if el:
                            break

            download = download_info.value
            original_name = download.suggested_filename or "resume.pdf"
            safe_name = re.sub(r'[\\/:*?"<>|]', '_', original_name)
            save_path = os.path.join(RESUME_DIR, safe_name)

            await download.save_as(save_path)
            resume_text = await self._extract_text(save_path)

            return {
                "status": "success",
                "file_path": save_path,
                "filename": safe_name,
                "file_size": os.path.getsize(save_path),
                "resume_text": resume_text[:3000] if resume_text else "",
            }
        except Exception as e:
            return {"status": "error", "message": f"点击下载失败: {str(e)}"}

    async def _extract_text(self, file_path: str) -> str:
        """从简历文件中提取文本内容."""
        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".pdf":
                return await self._extract_pdf_text(file_path)
            elif ext in (".doc", ".docx"):
                return await self._extract_docx_text(file_path)
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
        except Exception as e:
            log.warning(f"文本提取失败 ({ext}): {e}")
            return ""

    async def _extract_pdf_text(self, file_path: str) -> str:
        """提取PDF文本."""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            log.warning("pdfplumber 未安装，无法提取PDF文本")
            return ""
        except Exception as e:
            log.warning(f"PDF解析失败: {e}")
            return ""

    async def _extract_docx_text(self, file_path: str) -> str:
        """提取Word文档文本."""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts)
        except ImportError:
            log.warning("python-docx 未安装，无法提取Word文本")
            return ""
        except Exception as e:
            log.warning(f"Word解析失败: {e}")
            return ""

    async def batch_download_from_chat_list(self, max_count: int = 20) -> dict:
        """批量从聊天列表中下载所有有简历的对话.

        Args:
            max_count: 最多处理的对话数量
        Returns:
            批量下载结果
        """
        from chat_scraper import ChatScraper
        scraper = ChatScraper(self.browser)

        # 获取聊天列表
        chat_list = await scraper.get_chat_list()
        valid_chats = [c for c in chat_list if isinstance(c, dict) and "error" not in c]

        if not valid_chats:
            return {"error": "聊天列表为空"}

        all_results = []
        downloaded_count = 0

        for i, chat in enumerate(valid_chats[:max_count]):
            log.info(f"处理第 {i} 个对话: {chat.get('name', '未知')}")

            # 选择对话
            select_result = await scraper.select_chat_by_index(i)
            if "error" in select_result:
                continue

            # 检查是否有简历
            resume_check = await scraper.check_resume_received()
            if resume_check.get("has_resume"):
                # 下载简历
                download_result = await self.download_resume_from_chat()
                if download_result.get("downloaded", 0) > 0:
                    downloaded_count += download_result["downloaded"]
                    all_results.extend(download_result.get("details", []))

            await self.browser.random_delay()

        return {
            "status": "success",
            "total_chats": len(valid_chats),
            "processed": min(len(valid_chats), max_count),
            "resumes_downloaded": downloaded_count,
            "details": all_results,
        }
